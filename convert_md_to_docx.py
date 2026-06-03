from docx import Document
from pathlib import Path

in_path = Path('SS_Portal_Project_Report.md')
out_path = Path('SS_Portal_Project_Report2.docx')

if not in_path.exists():
    raise FileNotFoundError(f'Markdown file not found: {in_path}')

lines = in_path.read_text(encoding='utf-8').splitlines()

doc = Document()

for line in lines:
    stripped = line.strip()
    if stripped.startswith('### '):
        p = doc.add_paragraph(stripped[4:].strip(), style='Heading 3')
    elif stripped.startswith('## '):
        p = doc.add_paragraph(stripped[3:].strip(), style='Heading 2')
    elif stripped.startswith('# '):
        p = doc.add_paragraph(stripped[2:].strip(), style='Title')
    elif stripped.startswith('- [ ] '):
        p = doc.add_paragraph(stripped[6:].strip())
        p.style = 'List Bullet'
    elif stripped.startswith('- [x] '):
        p = doc.add_paragraph(stripped[6:].strip())
        p.style = 'List Bullet'
    elif stripped.startswith('- '):
        p = doc.add_paragraph(stripped[2:].strip())
        p.style = 'List Bullet'
    elif stripped == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

# Save file
doc.save(out_path)
print(f'Created {out_path.resolve()}')
